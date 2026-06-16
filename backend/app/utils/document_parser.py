import logging
from pathlib import Path
from typing import List
from langchain.schema import Document

logger = logging.getLogger(__name__)


def parse_document(file_path: str, filename: str) -> List[Document]:
    """
    Parse a document file into a list of LangChain Document objects.
    Supports PDF, DOCX, and TXT files.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_path, filename)
    elif ext == ".docx":
        return _parse_docx(file_path, filename)
    elif ext in [".txt", ".md"]:
        return _parse_text(file_path, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(file_path: str, filename: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    documents = []

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            documents.append(Document(
                page_content=text.strip(),
                metadata={"filename": filename, "page": page_num, "source": filename}
            ))

    logger.info(f"Parsed PDF '{filename}': {len(documents)} pages")
    return documents


def _parse_docx(file_path: str, filename: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    documents = [Document(
        page_content=full_text,
        metadata={"filename": filename, "page": 1, "source": filename}
    )]

    logger.info(f"Parsed DOCX '{filename}': {len(full_text)} characters")
    return documents


def _parse_text(file_path: str, filename: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    documents = [Document(
        page_content=text,
        metadata={"filename": filename, "page": 1, "source": filename}
    )]

    logger.info(f"Parsed TXT '{filename}': {len(text)} characters")
    return documents
